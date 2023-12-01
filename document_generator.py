import os
from docx import Document
from docx2pdf import convert


class DocumentGenerator:
    def __init__(self):
        pass

    def replace_fields(self, doc, row, field_mapping):
        for paragraph in doc.paragraphs:
            for field, col_name in field_mapping.items():
                for run in paragraph.runs:
                    run.text = run.text.replace(field, str(row[col_name]))

    def generate_document_for_row(
        self,
        row,
        template_doc,
        destination_folder,
        field_mapping,
        column_like_doc_name,
        file_name,
        pdf=False,
    ):
        if pdf:
            # Solo generar el documento PDF
            doc = Document(template_doc)
            self.replace_fields(doc, row, field_mapping)

            document_name = f"{file_name}_{str(row[field_mapping[column_like_doc_name]]).replace(' ', '_')}"

            # Guardar el documento docx temporal
            docx_path = os.path.join(destination_folder, f"{document_name}.docx")
            doc.save(docx_path)

            # Convertir el documento docx a PDF con convert()
            pdf_path = os.path.join(destination_folder, f"{document_name}.pdf")
            convert(docx_path, pdf_path)

            # Eliminar el documento docx temporal
            os.remove(docx_path)

            print(f"PDF generated for {document_name}. File saved at: {pdf_path}")
        else:
            # Generar solo el documento DOCX
            doc = Document(template_doc)
            self.replace_fields(doc, row, field_mapping)

            document_name = f"{file_name}_{str(row[field_mapping[column_like_doc_name]]).replace(' ', '_')}"
            new_document_path = os.path.join(
                destination_folder, f"{document_name}.docx"
            )
            doc.save(new_document_path)

            print(
                f"Document generated for {document_name}. File saved at: {new_document_path}"
            )
